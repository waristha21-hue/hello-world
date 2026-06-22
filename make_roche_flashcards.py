# -*- coding: utf-8 -*-
"""สร้างการ์ดสรุปสินค้าเรือธง Roche (Core Lab + Molecular) เป็นไฟล์ Word ปรินต์ได้"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Tahoma"  # มี glyph ไทยครบ ติดตั้งมากับ Windows ทุกเครื่อง

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="BBBBBB", sz="6"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), sz)
        e.set(qn('w:color'), color)
        borders.append(e)
    tcPr.append(borders)

def style_run(run, size=11, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    # ให้ font ไทยใช้ Tahoma ด้วย
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), FONT)
    rFonts.set(qn('w:hAnsi'), FONT)
    rFonts.set(qn('w:cs'), FONT)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def add_para(doc, text, size=11, bold=False, color=None, align=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    r = p.add_run(text)
    style_run(r, size=size, bold=bold, color=color)
    return p

def add_card(doc, title, header_fill, header_text_color, rows):
    """rows = list of (label, value)"""
    n = len(rows) + 1
    table = doc.add_table(rows=n, cols=2)
    table.autofit = False
    # ความกว้างคอลัมน์
    widths = [Cm(4.2), Cm(12.3)]
    # header (merge 2 ช่อง)
    hcells = table.rows[0].cells
    hcell = hcells[0].merge(hcells[1])
    set_cell_bg(hcell, header_fill)
    set_cell_borders(hcell)
    hp = hcell.paragraphs[0]
    hp.paragraph_format.space_after = Pt(2)
    hp.paragraph_format.space_before = Pt(2)
    hr = hp.add_run(title)
    style_run(hr, size=13, bold=True, color=header_text_color)
    # rows
    for i, (label, value) in enumerate(rows, start=1):
        c0, c1 = table.rows[i].cells
        set_cell_bg(c0, "F2F2F2")
        for c, w in zip((c0, c1), widths):
            c.width = w
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_borders(c)
        p0 = c0.paragraphs[0]; p0.paragraph_format.space_after = Pt(2); p0.paragraph_format.space_before = Pt(2)
        r0 = p0.add_run(label); style_run(r0, size=10.5, bold=True, color="333333")
        p1 = c1.paragraphs[0]; p1.paragraph_format.space_after = Pt(2); p1.paragraph_format.space_before = Pt(2)
        r1 = p1.add_run(value); style_run(r1, size=10.5)
    # บังคับความกว้างคอลัมน์
    for row in table.rows:
        for c, w in zip(row.cells, widths):
            c.width = w
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ---------- เริ่มสร้างเอกสาร ----------
doc = Document()
# ตั้ง font เริ่มต้น
style = doc.styles['Normal']
style.font.name = FONT
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:cs'), FONT)

for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

BLUE_FILL, BLUE_TXT = "1F6FB2", "FFFFFF"     # Core Lab
GREEN_FILL, GREEN_TXT = "38761D", "FFFFFF"   # Molecular

add_para(doc, "การ์ดสรุปสินค้าเรือธง Roche Diagnostics", size=18, bold=True, color="1F4E79", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "Core Lab + Molecular Lab  |  สำหรับเตรียมตัวเป็น Account Value Partner", size=11, color="666666", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

add_para(doc, "วิธีจำภาพรวมก่อน (กันสับสน)", size=13, bold=True, color="1F4E79", space_before=4, space_after=4)
add_para(doc, "•  cobas ตัวอักษร (pro / pure)  =  กลุ่ม Core Lab — ตรวจเลือดทั่วไป", size=11, space_after=2)
add_para(doc, "•  cobas ตัวเลข (5800 / 6800 / 8800)  =  กลุ่ม Molecular — ตรวจ DNA/RNA", size=11, space_after=2)
add_para(doc, "•  ยิ่งเลขเยอะ = ยิ่งปริมาณงานเยอะ = โรงพยาบาลยิ่งใหญ่  (5800 < 6800 < 8800)", size=11, space_after=10)

# การ์ด 1
add_card(doc, "การ์ด 1  —  cobas pro", BLUE_FILL, BLUE_TXT, [
    ("กลุ่ม", "Core Lab (เคมีคลินิก + ภูมิคุ้มกัน รวมในเครื่องเดียว)"),
    ("เหมาะกับ", "โรงพยาบาลใหญ่ / แล็บกลาง ปริมาณงานสูง"),
    ("จุดเด่น 1 ประโยค", "เครื่องตัวท็อป รวมทุกการตรวจเลือดหลักไว้ในระบบเดียว ปั่นงานได้เยอะ"),
    ("ตรวจอะไร", "น้ำตาล ไขมัน ตับ ไต + Elecsys (หัวใจ มะเร็ง ไทรอยด์ ติดเชื้อ)"),
    ("ประโยคขาย", "ลดจำนวนเครื่อง ลดคนเดินตัวอย่าง ผลออกเร็วขึ้น คุมงานปริมาณมากได้นิ่ง"),
    ("คู่แข่งตรงรุ่น", "Abbott Alinity ci  /  Siemens Atellica"),
])
# การ์ด 2
add_card(doc, "การ์ด 2  —  cobas pure", BLUE_FILL, BLUE_TXT, [
    ("กลุ่ม", "Core Lab (เคมี + ภูมิคุ้มกัน + เกลือแร่ ในเครื่องขนาดเล็ก)"),
    ("เหมาะกับ", "โรงพยาบาลกลาง-เล็ก / แล็บพื้นที่จำกัด"),
    ("จุดเด่น 1 ประโยค", "เครื่อง integrated ที่เล็กที่สุดของ Roche — ลงได้แม้พื้นที่น้อย"),
    ("ตรวจอะไร", "เหมือน cobas pro แต่ปริมาณงานน้อยกว่า (เมนูครบสำหรับ รพ.กลาง)"),
    ("ประโยคขาย", "ได้เทคโนโลยีระดับ รพ.ใหญ่ ในขนาดที่ รพ.เล็กลงได้ ประหยัดพื้นที่+คน"),
    ("คู่แข่งตรงรุ่น", "Abbott Alinity (รุ่นเล็ก)  /  Beckman"),
])
# การ์ด 3
add_card(doc, "การ์ด 3  —  cobas 5800", GREEN_FILL, GREEN_TXT, [
    ("กลุ่ม", "Molecular (PCR ตรวจ DNA/RNA)"),
    ("เหมาะกับ", "โรงพยาบาลกลาง / แล็บเริ่มต้นงานโมเลกุล"),
    ("จุดเด่น 1 ประโยค", "ประตูสู่งานโมเลกุล — เครื่องเล็ก ใส่ตัวอย่างแล้วเดินจากได้ (sample-to-result)"),
    ("ตรวจอะไร", "Viral Load (HIV/HBV/HCV), HPV, CT/NG"),
    ("ประโยคขาย", "เริ่มทำ molecular เองได้ ไม่ต้องส่งตรวจข้างนอก คุมคุณภาพ + ลดเวลารอผล"),
    ("คู่แข่งตรงรุ่น", "Cepheid GeneXpert  /  Abbott Alinity m (รุ่นเล็ก)"),
])
# การ์ด 4
add_card(doc, "การ์ด 4  —  cobas 6800", GREEN_FILL, GREEN_TXT, [
    ("กลุ่ม", "Molecular (PCR อัตโนมัติเต็มรูปแบบ)"),
    ("เหมาะกับ", "โรงพยาบาลใหญ่ / ศูนย์ตรวจปริมาณงานสูง"),
    ("จุดเด่น 1 ประโยค", "เครื่องโมเลกุลอัตโนมัติยอดนิยม — เดินงานเองได้ทั้งกะ (walk-away)"),
    ("ปริมาณงาน", "~96 ตัวอย่าง/รอบ, ได้ผลถึง ~384 ผล/กะ 8 ชม. (เลขโดยประมาณ)"),
    ("ประโยคขาย", "ลดแรงคน ทำงานทั้งกะแบบไม่ต้องเฝ้า รองรับงานพีคได้สบาย"),
    ("คู่แข่งตรงรุ่น", "Hologic Panther  /  Abbott Alinity m"),
])
# การ์ด 5
add_card(doc, "การ์ด 5  —  cobas 8800", GREEN_FILL, GREEN_TXT, [
    ("กลุ่ม", "Molecular (ปริมาณงานสูงสุดของ Roche)"),
    ("เหมาะกับ", "แล็บอ้างอิง / ศูนย์ตรวจระดับชาติ ปริมาณงานมหาศาล"),
    ("จุดเด่น 1 ประโยค", "พี่ใหญ่สุดของตระกูล — ปั่นงานโมเลกุลได้มากที่สุด"),
    ("ปริมาณงาน", "~282 ตัวอย่าง/รอบ, ได้ผลถึง ~960 ผล/กะ 8 ชม. (เลขโดยประมาณ)"),
    ("ประโยคขาย", "สเกลใหญ่สุด รองรับงานคัดกรองระดับประเทศ ต้นทุนต่อ test ต่ำเมื่องานเยอะ"),
    ("คู่แข่งตรงรุ่น", "Hologic Panther  /  Abbott Alinity m (รุ่นใหญ่)"),
])

# เทคนิคท่องจำ
add_para(doc, "เทคนิคท่องจำเร็ว", size=13, bold=True, color="1F4E79", space_before=6, space_after=4)
add_para(doc, "•  บันได Molecular:  5800 (เริ่มต้น) → 6800 (มาตรฐาน รพ.ใหญ่) → 8800 (ใหญ่สุด) — ยิ่งเลขสูง ยิ่งงานเยอะ", size=11, space_after=2)
add_para(doc, "•  คู่ Core Lab:  pro = “โปร” ใหญ่  /  pure = “เพียวร์” เล็กกะทัดรัด", size=11, space_after=2)
add_para(doc, "•  ตัวเลขปริมาณงานของ 6800/8800 เป็นค่าโดยประมาณ — ตอนเข้างานจริงให้เช็กตัวเลขเป๊ะจาก spec sheet ภายในของ Roche อีกครั้ง", size=10.5, color="C0392B", space_after=2)

from pathlib import Path
out = Path(__file__).parent / "Roche_Flashcards.docx"
doc.save(out)
print(f"SAVED: {out}")
